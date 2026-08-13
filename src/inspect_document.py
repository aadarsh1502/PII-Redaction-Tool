from docx import Document
from pathlib import Path


INPUT_FILE = Path("input/Red Herring Prospectus.docx")


def inspect_document():
    doc = Document(INPUT_FILE)

    print("=" * 60)
    print("DOCUMENT INSPECTION")
    print("=" * 60)

    # Basic information
    print(f"File: {INPUT_FILE}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    print(f"Sections: {len(doc.sections)}")

    # Paragraph information
    non_empty_paragraphs = [
        p for p in doc.paragraphs
        if p.text.strip()
    ]

    print(f"Non-empty paragraphs: {len(non_empty_paragraphs)}")

    # Tables
    total_rows = 0
    total_cells = 0

    for table in doc.tables:
        total_rows += len(table.rows)
        for row in table.rows:
            total_cells += len(row.cells)

    print(f"Total table rows: {total_rows}")
    print(f"Total table cells: {total_cells}")

    # Headers and footers
    header_count = 0
    footer_count = 0

    for section in doc.sections:
        header_count += sum(
            1 for p in section.header.paragraphs
            if p.text.strip()
        )

        footer_count += sum(
            1 for p in section.footer.paragraphs
            if p.text.strip()
        )

    print(f"Non-empty header paragraphs: {header_count}")
    print(f"Non-empty footer paragraphs: {footer_count}")

    print("=" * 60)

    # Show a few sample paragraphs
    print("\nSAMPLE PARAGRAPHS")
    print("=" * 60)

    for i, paragraph in enumerate(non_empty_paragraphs[:10]):
        print(f"\n[{i}]")
        print(paragraph.text[:500])


if __name__ == "__main__":
    inspect_document()