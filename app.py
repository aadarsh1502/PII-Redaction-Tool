import os
import sys
from io import BytesIO

import spacy
import streamlit as st
from docx import Document
from faker import Faker

# ============================================================
# EXISTING BACKEND — DO NOT MODIFY
# ============================================================

ROOT_DIR = os.path.dirname(os.path.abspath(__file__))
SRC_DIR = os.path.join(ROOT_DIR, "src")

if SRC_DIR not in sys.path:
    sys.path.insert(0, SRC_DIR)

from detectors import detect_all_pii


# ============================================================
# FAKER
# ============================================================

fake = Faker("en_IN")


# ============================================================
# LOAD SPACY MODEL
# ============================================================

@st.cache_resource
def load_nlp():
    return spacy.load("en_core_web_sm")


# ============================================================
# GENERATE FAKE VALUES
# ============================================================

def generate_fake_value(pii_type):
    if pii_type == "PERSON":
        return fake.name()

    if pii_type == "COMPANY":
        return fake.company()

    if pii_type == "ADDRESS":
        return fake.address().replace("\n", ", ")

    if pii_type == "EMAIL":
        return fake.email()

    if pii_type == "PHONE":
        return fake.phone_number()

    if pii_type == "DOB":
        return fake.date_of_birth().strftime("%d-%m-%Y")

    if pii_type == "SSN":
        return fake.bothify(text="#########")

    if pii_type == "CREDIT_CARD":
        return fake.credit_card_number()

    if pii_type == "IP_ADDRESS":
        return fake.ipv4()

    return "[REDACTED]"


# ============================================================
# REDACT TEXT
# ============================================================

def redact_text(text, nlp, pii_mapping):
    if not text.strip():
        return text, []

    results = detect_all_pii(
        text,
        nlp=nlp
    )

    if not results:
        return text, []

    redacted_text = text

    # Replace from right to left so character positions
    # remain valid after each replacement.
    for result in sorted(
        results,
        key=lambda item: item["start"],
        reverse=True
    ):
        pii_type = result["type"]
        original_value = result["value"]

        # Use the original value as the mapping key.
        # This keeps repeated PII consistent throughout
        # the entire document.
        mapping_key = (
            pii_type,
            original_value.strip()
        )

        if mapping_key not in pii_mapping:
            pii_mapping[mapping_key] = generate_fake_value(
                pii_type
            )

        replacement = pii_mapping[mapping_key]

        start = result["start"]
        end = result["end"]

        redacted_text = (
            redacted_text[:start]
            + replacement
            + redacted_text[end:]
        )

    return redacted_text, results


# ============================================================
# PROCESS DOCX
# ============================================================

def process_docx(file_bytes, nlp):
    doc = Document(BytesIO(file_bytes))

    counts = {}
    total = 0

    # Stores original PII -> fake PII.
    # This ensures repeated values get the same replacement.
    pii_mapping = {}

    # ========================================================
    # PROCESS PARAGRAPHS
    # ========================================================

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue

        redacted_text, results = redact_text(
            paragraph.text,
            nlp,
            pii_mapping
        )

        if results:
            paragraph.text = redacted_text

            for result in results:
                pii_type = result["type"]

                counts[pii_type] = (
                    counts.get(pii_type, 0) + 1
                )

                total += 1

    # ========================================================
    # PROCESS TABLES
    # ========================================================

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        continue

                    redacted_text, results = redact_text(
                        paragraph.text,
                        nlp,
                        pii_mapping
                    )

                    if results:
                        paragraph.text = redacted_text

                        for result in results:
                            pii_type = result["type"]

                            counts[pii_type] = (
                                counts.get(pii_type, 0) + 1
                            )

                            total += 1

    # ========================================================
    # SAVE DOCUMENT TO MEMORY
    # ========================================================

    output = BytesIO()

    doc.save(output)

    output.seek(0)

    return (
        output.getvalue(),
        counts,
        total,
        pii_mapping
    )


# ============================================================
# STREAMLIT UI
# ============================================================

st.set_page_config(
    page_title="PII Redactor",
    page_icon="🔒",
    layout="centered"
)

st.title("🔒 PII Redactor")

st.write(
    "Upload a Microsoft Word document to detect and "
    "anonymize personally identifiable information."
)

st.info(
    "Detected PII is replaced with realistic synthetic "
    "values. Your original document is never modified."
)


# ============================================================
# FILE UPLOAD
# ============================================================

uploaded_file = st.file_uploader(
    "Upload a DOCX file",
    type=["docx"]
)


if uploaded_file is not None:
    st.success(
        f"Uploaded: {uploaded_file.name}"
    )

    if st.button(
        "🔍 Scan & Anonymize",
        type="primary",
        use_container_width=True
    ):
        with st.spinner(
            "Detecting and anonymizing PII..."
        ):
            try:
                nlp = load_nlp()

                (
                    redacted_bytes,
                    counts,
                    total,
                    pii_mapping
                ) = process_docx(
                    uploaded_file.getvalue(),
                    nlp
                )

                st.success(
                    "Document successfully anonymized!"
                )

                # =================================================
                # PII SUMMARY
                # =================================================

                st.subheader(
                    "PII Detection Summary"
                )

                if counts:
                    for pii_type in sorted(counts):
                        st.write(
                            f"**{pii_type}:** "
                            f"{counts[pii_type]}"
                        )

                    st.divider()

                    st.metric(
                        "Total PII Detected",
                        total
                    )

                else:
                    st.info(
                        "No PII was detected."
                    )

                # =================================================
                # ANONYMIZATION PREVIEW
                # =================================================

                if pii_mapping:
                    st.subheader(
                        "Anonymization Preview"
                    )

                    preview = []

                    for (
                        (pii_type, original),
                        replacement
                    ) in pii_mapping.items():
                        preview.append(
                            {
                                "Type": pii_type,
                                "Original": original,
                                "Replacement": replacement
                            }
                        )

                    st.dataframe(
                        preview,
                        use_container_width=True,
                        hide_index=True
                    )

                # =================================================
                # DOWNLOAD
                # =================================================

                base_name = os.path.splitext(
                    uploaded_file.name
                )[0]

                output_name = (
                    f"Anonymized_{base_name}.docx"
                )

                st.download_button(
                    label="⬇️ Download Anonymized Document",
                    data=redacted_bytes,
                    file_name=output_name,
                    mime=(
                        "application/vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"
                    ),
                    use_container_width=True
                )

            except Exception as e:
                st.error(
                    "An error occurred while processing "
                    "the document."
                )

                st.exception(e)