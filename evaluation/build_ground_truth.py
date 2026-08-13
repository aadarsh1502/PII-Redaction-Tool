"""
build_ground_truth.py
---------------------
Utility script that prints a representative sample of paragraphs and table
cells from the source document together with the PII detected by
detect_all_pii.  The output was used to manually curate ground_truth.json.

Run from the project root:
    python evaluation/build_ground_truth.py
"""

import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

import spacy
from docx import Document
from detectors import detect_all_pii

INPUT_FILE = os.path.join(
    os.path.dirname(__file__), '..', 'input', 'Red Herring Prospectus.docx'
)


def main():
    nlp = spacy.load('en_core_web_sm')
    doc = Document(INPUT_FILE)

    print('=' * 70)
    print('PARAGRAPH PII SAMPLE')
    print('=' * 70)

    shown = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        results = detect_all_pii(text, nlp)
        if results:
            print(f'\n[P{i}] {text[:300]}')
            for r in results:
                print(f'  [{r["type"]}] {repr(r["value"])}')
            shown += 1
        if shown >= 15:
            break

    print('\n' + '=' * 70)
    print('TABLE CELL PII SAMPLE')
    print('=' * 70)

    shown = 0
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue
                results = detect_all_pii(text, nlp)
                if results:
                    print(f'\n[T{ti}R{ri}C{ci}] {text[:300]}')
                    for r in results:
                        print(f'  [{r["type"]}] {repr(r["value"])}')
                    shown += 1
            if shown >= 15:
                break
        if shown >= 15:
            break


if __name__ == '__main__':
    main()
