import spacy
from docx import Document

from detectors import detect_all_pii


INPUT_FILE = "input/Red Herring Prospectus.docx"


def extract_text(doc):
    parts = []

    # Extract paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def main():

    # --------------------------------------------------------
    # Load NER model
    # --------------------------------------------------------
    nlp = spacy.load("en_core_web_sm", exclude=["parser", "tagger", "lemmatizer", "attribute_ruler"])

    # --------------------------------------------------------
    # Load document
    # --------------------------------------------------------
    doc = Document(INPUT_FILE)
    text = extract_text(doc)

    # --------------------------------------------------------
    # Run complete PII detection pipeline
    # --------------------------------------------------------
    results = detect_all_pii(
        text,
        nlp
    )

    # --------------------------------------------------------
    # Print results
    # --------------------------------------------------------
    print("=" * 70)
    print("PII SCAN OF PROSPECTUS")
    print("=" * 70)

    counts = {}

    for result in results:

        pii_type = result["type"]

        print(
            f"[{pii_type}] {result['value']}"
        )

        counts[pii_type] = (
            counts.get(pii_type, 0) + 1
        )

    # --------------------------------------------------------
    # Summary
    # --------------------------------------------------------
    print("=" * 70)

    for pii_type, count in sorted(counts.items()):

        print(
            f"{pii_type} entities: {count}"
        )

    print(
        f"Total PII entities: {len(results)}"
    )

    print("=" * 70)


if __name__ == "__main__":
    main()