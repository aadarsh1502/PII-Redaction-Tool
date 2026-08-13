import os
import re
import spacy
from docx import Document
from faker import Faker

from detectors import detect_all_pii, CompanyNameRegistry


INPUT_FILE = "input/Red Herring Prospectus.docx"
OUTPUT_FILE = "output/Redacted_Prospectus.docx"

_faker = Faker()
Faker.seed(42)


# ============================================================
# FAKER-BASED REDACTOR
# ============================================================

class FakerRedactor:
    """
    Generates realistic fake replacements per PII type and caches them so
    the same original value always maps to the same fake value within a
    single document run.

    Cache key: (pii_type, original_value)
    """

    def __init__(self):
        self._mapping: dict[tuple[str, str], str] = {}

    # ----------------------------------------------------------
    # Internal generators
    # ----------------------------------------------------------

    def _generate(self, pii_type: str, original_value: str) -> str:

        if pii_type == "PERSON":
            return _faker.name()

        if pii_type == "COMPANY":
            return _faker.company()

        if pii_type == "EMAIL":
            return _faker.email()

        if pii_type == "PHONE":
            # Preserve country-code prefix (+XX or +XXX) if present in original
            cc_match = re.match(r"(\+\d{1,3})", original_value.strip())
            if cc_match:
                cc = cc_match.group(1)
                # Generate a local number (10 digits) and attach the same CC
                local = _faker.numerify("##########")
                return f"{cc} {local[:5]} {local[5:]}"
            return _faker.phone_number()

        if pii_type == "ADDRESS":
            return _faker.address().replace("\n", ", ")

        if pii_type == "DOB":
            return _faker.date_of_birth(
                minimum_age=18, maximum_age=80
            ).strftime("%d/%m/%Y")

        if pii_type == "SSN":
            return _faker.ssn()

        if pii_type == "CREDIT_CARD":
            return _faker.credit_card_number(card_type=None)

        if pii_type == "IP_ADDRESS":
            return _faker.ipv4()

        # Fallback
        return f"[REDACTED {pii_type}]"

    # ----------------------------------------------------------
    # Public interface
    # ----------------------------------------------------------

    def get_fake(self, pii_type: str, original_value: str) -> str:
        """
        Return the fake replacement for (pii_type, original_value).
        Generates and caches a new value if not seen before.

        NOTE: for COMPANY detections that came from the shortened-
        mention registry (see CompanyNameRegistry), the caller should
        pass the CANONICAL full company name as original_value so that
        "ICICI Securities" and "ICICI Securities Limited" both resolve
        to the exact same fake company name. This is handled in
        redact_text() below via the "canonical" field on the detection.
        """
        key = (pii_type, original_value)
        if key not in self._mapping:
            self._mapping[key] = self._generate(pii_type, original_value)
        return self._mapping[key]


# ============================================================
# REDACT TEXT
# ============================================================

def redact_text(text: str, nlp, redactor: FakerRedactor, company_registry: CompanyNameRegistry):
    """
    Detect all PII in *text*, replace each span with a fake value sourced
    from *redactor*, and return the redacted string plus the list of
    detection dicts.
    """

    if not text.strip():
        return text, []

    results = detect_all_pii(text, nlp=nlp, company_registry=company_registry)

    if not results:
        return text, []

    # Replace from right-to-left so earlier character positions stay valid.
    redacted_text = text

    for result in sorted(results, key=lambda item: item["start"], reverse=True):

        # For shortened company mentions (e.g. "ICICI Securities" found
        # via the registry), map using the CANONICAL full name so the
        # fake value stays consistent with the full-suffix mention
        # elsewhere in the document.
        mapping_key_value = result.get("canonical", result["value"])

        fake_value = redactor.get_fake(result["type"], mapping_key_value)

        start = result["start"]
        end = result["end"]

        redacted_text = (
            redacted_text[:start]
            + fake_value
            + redacted_text[end:]
        )

    return redacted_text, results


# ============================================================
# BUILD COMPANY REGISTRY (whole-document warm-up pass)
# ============================================================

def build_company_registry(doc, nlp) -> CompanyNameRegistry:
    """
    Scan the ENTIRE document once (paragraphs + tables) purely to collect
    every company name that validates with its full legal suffix (e.g.
    "ICICI Securities Limited"). This lets the main redaction pass catch
    shortened/informal repeat mentions ("ICICI Securities") consistently,
    even if the shortened form appears BEFORE the full-suffix mention.
    """
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)

    full_document_text = "\n".join(text_parts)

    registry = CompanyNameRegistry()
    registry.warm_up(full_document_text, nlp)

    return registry


# ============================================================
# PROCESS DOCUMENT
# ============================================================

def process_document():

    print("=" * 70)
    print("PII REDACTION")
    print("=" * 70)

    # Load spaCy model (NER only — parser/tagger/etc excluded to avoid
    # memory issues on repeated per-paragraph calls)
    nlp = spacy.load(
        "en_core_web_sm",
        exclude=["parser", "tagger", "lemmatizer", "attribute_ruler"],
    )

    # Load document
    doc = Document(INPUT_FILE)

    # One redactor instance per document run — shared mapping across all
    # paragraphs, tables, headers, and footers.
    redactor = FakerRedactor()

    # Warm-up pass: scan the whole document once to build the company
    # name registry BEFORE doing any redaction, so shortened company
    # mentions are caught consistently regardless of document order.
    print("Building company-name registry (whole-document warm-up pass)...")
    company_registry = build_company_registry(doc, nlp)
    print(f"Registry built: {len(company_registry._canonical)} canonical company name(s) found.\n")

    total_detections = 0
    counts = {}

    # --------------------------------------------------------
    # Helper: process a single paragraph in-place
    # --------------------------------------------------------

    def _process_paragraph(paragraph):
        nonlocal total_detections

        if not paragraph.text.strip():
            return

        original_text = paragraph.text
        redacted_text, results = redact_text(
            original_text, nlp, redactor, company_registry
        )

        if results:
            paragraph.text = redacted_text
            total_detections += len(results)
            for result in results:
                pii_type = result["type"]
                counts[pii_type] = counts.get(pii_type, 0) + 1

    # --------------------------------------------------------
    # Paragraphs
    # --------------------------------------------------------

    for paragraph in doc.paragraphs:
        _process_paragraph(paragraph)

    # --------------------------------------------------------
    # Tables
    # --------------------------------------------------------

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _process_paragraph(paragraph)

    # --------------------------------------------------------
    # Headers & Footers (every section)
    # --------------------------------------------------------

    hf_detections = 0

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            before = total_detections
            _process_paragraph(paragraph)
            hf_detections += total_detections - before

        for paragraph in section.footer.paragraphs:
            before = total_detections
            _process_paragraph(paragraph)
            hf_detections += total_detections - before

    # --------------------------------------------------------
    # Create output directory
    # --------------------------------------------------------

    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)

    # Save — fall back to a timestamped name if the canonical path is
    # locked by another process (e.g. the file is open in Word).
    actual_output = OUTPUT_FILE
    try:
        doc.save(OUTPUT_FILE)
    except PermissionError:
        import datetime
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        base, ext = os.path.splitext(OUTPUT_FILE)
        actual_output = f"{base}_{ts}{ext}"
        doc.save(actual_output)
        print(
            f"\n[WARNING] {OUTPUT_FILE} is locked by another process.\n"
            f"          Saved to {actual_output} instead.\n"
        )

    # --------------------------------------------------------
    # Summary
    # --------------------------------------------------------

    print("\nREDACTION SUMMARY")
    print("=" * 70)

    for pii_type in sorted(counts):
        print(f"{pii_type:<15}: {counts[pii_type]}")

    print("-" * 70)

    print(
        f"HEADER/FOOTER   : {hf_detections}"
    )

    print(
        f"TOTAL REDACTIONS: {total_detections}"
    )

    print("-" * 70)

    print(f"OUTPUT FILE     : {actual_output}")

    print("=" * 70)


# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":
    process_document()
