import spacy
from docx import Document

from detectors import detect_all_pii


INPUT_FILE = "input/Red Herring Prospectus.docx"


def scan_text(text, nlp):
    """Run the exact same detector pipeline used by redaction."""
    if not text or not text.strip():
        return []

    return detect_all_pii(text, nlp)


def main():

    print("=" * 70)
    print("FULL PII SCAN OF PROSPECTUS")
    print("=" * 70)

    nlp = spacy.load("en_core_web_sm", exclude=["parser", "tagger", "lemmatizer", "attribute_ruler"])
    doc = Document(INPUT_FILE)

    counts = {
        "ADDRESS": 0,
        "COMPANY": 0,
        "EMAIL": 0,
        "PERSON": 0,
        "PHONE": 0,
        "DOB": 0,
        "SSN": 0,
        "CREDIT_CARD": 0,
        "IP_ADDRESS": 0,
    }

    total = 0

    # --------------------------------------------------------
    # Paragraphs
    # --------------------------------------------------------

    for paragraph in doc.paragraphs:

        text = paragraph.text

        if not text.strip():
            continue

        results = scan_text(text, nlp)

        for result in results:

            pii_type = result["type"]

            if pii_type in counts:
                counts[pii_type] += 1

            total += 1

            print(
                f"[{pii_type}] "
                f"{result['value']}"
            )

    # --------------------------------------------------------
    # Tables
    # --------------------------------------------------------

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    text = paragraph.text

                    if not text.strip():
                        continue

                    results = scan_text(
                        text,
                        nlp
                    )

                    for result in results:

                        pii_type = result["type"]

                        if pii_type in counts:
                            counts[pii_type] += 1

                        total += 1

                        print(
                            f"[{pii_type}] "
                            f"{result['value']}"
                        )

    # --------------------------------------------------------
    # Summary
    # --------------------------------------------------------

    print("=" * 70)
    print("PII DETECTION SUMMARY")
    print("=" * 70)

    for pii_type, count in counts.items():

        print(
            f"{pii_type:<15}: {count}"
        )

    print("-" * 70)

    print(
        f"TOTAL PII       : {total}"
    )

    print("=" * 70)


if __name__ == "__main__":
    main()